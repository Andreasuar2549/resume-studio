#!/usr/bin/env python3
"""
build_docx.py — render a resume.md master file into an ATS-safe .docx.

ATS-safe by construction: single column, no tables/text boxes/headers/footers,
standard fonts, real bullet paragraphs, all content in the document body.

Markdown conventions expected in resume.md:

    # Full Name
    email · phone · city · linkedin.com/in/x · github.com/x   <- contact line(s), before first ##

    ## Summary
    - bullet or plain paragraph(s)

    ## Work Experience
    ### Senior Product Manager | Acme Corp | Bangalore | Mar 2022 – Present
    - Achievement bullet...

    (H3 pipe fields are flexible: "Title | Company | Dates" also works — the
     LAST pipe field is treated as the right-aligned date if it contains a digit
     or "Present"; everything else joins the bold left side.)

    ## Education
    ### B.Tech, Computer Science | IIT Delhi | 2016 – 2020

    ## Skills
    Plain paragraphs or bullets.

Inline **bold** and *italic* are supported. Everything else is plain text.

Usage:
    python build_docx.py resume.md [-o resume.docx] [--font Calibri] [--size 10.5]
                                   [--margin 0.6] [--a4]
Then convert to PDF if needed:  soffice --headless --convert-to pdf resume.docx
"""
import argparse
import re
import sys

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
except ImportError:
    sys.exit("python-docx is required: pip install python-docx --break-system-packages")

BLACK = RGBColor(0, 0, 0)
INLINE = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*)")


def add_runs(par, text, base_bold=False):
    """Write text into paragraph, honoring **bold** / *italic* spans."""
    for chunk in INLINE.split(text):
        if not chunk:
            continue
        bold, italic = base_bold, False
        if chunk.startswith("**") and chunk.endswith("**"):
            bold, chunk = True, chunk[2:-2]
        elif chunk.startswith("*") and chunk.endswith("*"):
            italic, chunk = True, chunk[1:-1]
        run = par.add_run(chunk)
        run.bold, run.italic = bold, italic
        run.font.color.rgb = BLACK


def looks_like_date(field):
    f = field.strip().lower()
    return bool(re.search(r"\d", f)) or "present" in f


def build(md_path, out_path, font, size, margin, a4):
    lines = open(md_path, encoding="utf-8").read().splitlines()
    doc = Document()

    sec = doc.sections[0]
    if a4:
        sec.page_width, sec.page_height = Inches(8.27), Inches(11.69)
    for side in ("top", "bottom", "left", "right"):
        setattr(sec, f"{side}_margin", Inches(margin))

    normal = doc.styles["Normal"]
    normal.font.name, normal.font.size = font, Pt(size)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.space_before = Pt(0)

    usable_width = (sec.page_width - sec.left_margin - sec.right_margin)
    seen_section = False

    def para(style=None):
        return doc.add_paragraph(style=style)

    for raw in lines:
        line = raw.rstrip()
        if not line.strip():
            continue
        if line.lstrip().startswith("<!--"):
            continue

        if line.startswith("# ") and not line.startswith("## "):          # Name
            p = para()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(line[2:].strip())
            r.bold, r.font.size = True, Pt(size + 6)
            r.font.color.rgb = BLACK
            p.paragraph_format.space_after = Pt(2)

        elif line.startswith("## "):                                       # Section
            seen_section = True
            p = para()
            r = p.add_run(line[3:].strip().upper())
            r.bold, r.font.size = True, Pt(size + 1)
            r.font.color.rgb = BLACK
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(3)

        elif line.startswith("### "):                                      # Role / degree
            fields = [f.strip() for f in line[4:].split("|")]
            date = fields.pop() if len(fields) > 1 and looks_like_date(fields[-1]) else None
            p = para()
            p.paragraph_format.space_before = Pt(5)
            add_runs(p, ", ".join(fields), base_bold=True)
            if date:
                p.paragraph_format.tab_stops.add_tab_stop(usable_width, WD_TAB_ALIGNMENT.RIGHT)
                r = p.add_run("\t" + date)
                r.italic = False
                r.font.color.rgb = BLACK

        elif re.match(r"^\s*[-*•]\s+", line):                              # Bullet
            text = re.sub(r"^\s*[-*•]\s+", "", line)
            p = para(style="List Bullet")
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Inches(0.22)
            add_runs(p, text)

        else:                                                              # Plain text
            p = para()
            if not seen_section:                                           # contact line(s)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(2)
            add_runs(p, line.strip())

    doc.save(out_path)
    words = sum(len(l.split()) for l in lines if not l.startswith("#"))
    print(f"Wrote {out_path}  (~{words} words in body — playbook target is 475-600)")


if __name__ == "__main__":
    ap = argparse.ArgumentParser(description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("input")
    ap.add_argument("-o", "--output", default=None)
    ap.add_argument("--font", default="Calibri")
    ap.add_argument("--size", type=float, default=10.5)
    ap.add_argument("--margin", type=float, default=0.6, help="inches, all sides")
    ap.add_argument("--a4", action="store_true", help="A4 page (default US Letter)")
    a = ap.parse_args()
    build(a.input, a.output or a.input.rsplit(".", 1)[0] + ".docx", a.font, a.size, a.margin, a.a4)
